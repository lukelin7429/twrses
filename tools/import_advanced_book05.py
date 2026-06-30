#!/usr/bin/env python3
import json
import re
from collections import Counter
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
BOOK = 5
SOURCE_DIR = Path.home() / "Downloads" / "高級閱讀 Book05"
DATA_PATH = ROOT / "data" / "advanced.json"
IMG_DIR = ROOT / "assets" / "img" / "booklets"

UNIT_FILES = {
    1: "Copy of Advanced_Book05_Unit01.doc.docx",
    2: "Copy of Advanced_Book05_Unit02.doc.docx",
    3: "Copy of Advanced_Book05_Unit03.doc.docx",
    4: "Copy of Advanced_Book05_Unit04.doc.docx",
    5: "Copy of Advanced_Book05_Unit05.doc.docx",
}

TRANSLATIONS = {
    1: [
        "竹塘是彰化縣北斗附近一個寧靜的小村莊，也是我的好友、同事兼醫療顧問 Luke 的家鄉。在這裡，你不會看到壅塞的道路、擁擠的市場，或主宰天際線的摩天大樓。你會看到的是許多開闊的田地、看似空蕩的街道，以及不會忙到錯過和陌生人交談機會的友善臉孔。這裡並不起眼，但在簡單的小建築與平坦濕潤的稻田之間，有一樣東西特別突出，那就是明航寺。",
        "竹塘似乎不是一座宏偉寺院會出現的地方，但它確實矗立在這裡。以它服務的人口稀少社區來看，寺院的規模顯得格外龐大。不過，這座宗教建築有一點和竹塘非常契合：安靜，以及毫無阻礙的簡樸。",
        "我對這座寺院最先注意到的是設計上的簡潔。台灣多數寺廟似乎會以大量象徵讓人目不暇給，裡面一幅接一幅的畫、一座接一座的雕塑，都反映出數千年的文化。明航寺很不一樣。極簡的風格從大門開始迎接你，不會讓你困惑，也不企圖支配或壓倒你；但仍然有足夠可看的地方，而且建築風格無疑帶有亞洲特色。",
        "庭院非常出色。修剪整齊的灌木、樹木與植物，填滿了鋪著頁岩的花園。兩座厚實的大象雕像從大門向外望著，彷彿守護著這片聖地。這個空間真正預示了人們在平和寺院內會感受到的寧靜。我可以想像自己在這樣的庭院裡，在樹下休息，讀著一個關於久遠浪漫文明的美好故事，並感到十分滿足。",
        "建築本身和前方的大門一樣非常對稱。中央是一座大型主體建築，兩側各有較小的建築。這裡的樸素更加凸顯建築的實用性，也透露出居住於此的人們所具有的真誠與樸實。真正的佛殿位於建築中心，空間寬闊，兩側有大型門廳通往建築其他部分。住在修道院裡的尼師照料著這座龐大的地方，這一定是許多工作。光是庭園維護就會是一場持續不斷的工作，更不用說主建築內每天都需要清潔。但對於把生命奉獻給佛陀的人來說，這一定是一份出於熱愛的勞動。",
        "我很幸運，在參訪寺院時受邀和尼師們共進午餐。她們帶著大大的笑容迎接我，雙手合十的樣子有點像基督徒祈禱。素食料理以類似自助餐的方式擺在長桌上。Luke 遞給我一個碗，要我盡情享用。尼師們準備了各式各樣我從未吃過的美味食物。我把盤子裝滿異國風味的佳餚，坐下來面對隔壁桌的兩位尼師。她們在我吃飯時偶爾看過來，點頭表示認可，也許是對我使用筷子的能力感到滿意。一位和善的女士給了我一些花生，另一位給了我一些我吃過最好的香蕉。用餐時，庭院裡的鳥兒用美麗的歌聲為我們演奏，那是在大城市裡永遠聽不到的聲音。它似乎正好成為這頓飯最合適的伴奏。",
        "午餐後，Luke 帶我簡短參觀建築的其他部分。他讓我看他用來教授免費英語課的教室。他希望藉由教當地孩子英文，有一天他們能把佛教智慧傳播給其他國家的人。",
        "我非常感謝有機會參訪明航寺。我不會很快忘記在那裡感受到的平靜，也不會忘記尼師們對我這個完全陌生的人所展現的款待。雖然我有時被認為有點悲觀，尤其是對人性，但這個美好的地方讓我看到世界上仍然有善。希望這座寺院，以及其他類似的地方，能持續滋養並擴展我們所有人內在與外在的善。",
    ],
    2: [
        "湖東國小的一些學生正在進行一項大型藝術計畫。他們在學校不同區域繪製幾幅大型壁畫。完成後，這些壁畫真的會讓學校看起來漂亮許多。",
        "這項計畫的主題是台灣的甘蔗產業。有些畫作已經完成，呈現不同場景，例如甘蔗的採收與運送。甘蔗在台灣發展史上扮演重要角色，因此這個產業被選為計畫主題。",
        "第一幅完成的畫，是一戶農家在田裡工作的景象。他們忙著砍下周圍茂盛的甘蔗作物。真人大小的描繪真的讓人覺得自己就在現場，與人們一起辛苦工作。計畫進行到一定程度後，開始受到許多關注。家長和社區居民特地到學校來看壁畫。人們對作品的巨大規模感到驚訝，也對作品品質感到驚喜。",
        "第二幅畫似乎特別吸引人。它描繪一輛老式蒸汽火車，拖著一車甘蔗穿越具有台灣特色的風景，令人印象深刻。前景中的一棵孤樹為畫作增添了許多深度，有些人甚至伸手去摸畫，彷彿以為它是立體的。",
        "第三幅畫的初步工作目前正在進行。當地美術老師，也是這項計畫幕後推手的紀棋元老師，已經粗略畫出基本形狀與設計。這幅畫會是一頭牛拉著載滿甘蔗的車。和其他壁畫一樣，學生們會在紀老師專業的指導下完成畫面。不過，在這項計畫上已經投入這麼多小時之後，大多數學生只需要偶爾一點建議。",
        "紀老師的教學真的幫助學生學會調色。在第一幅壁畫中，顏料必須一再混合、重新調整，直到顏色正確為止。但到了第三幅畫時，學生們已經抓到訣竅。這項雄心勃勃的工作會讓每個人受益。參與的學生獲得製作大型繪畫的寶貴經驗，學校也得到迫切需要的美化。也許最大的好處，是讓台灣豐富多彩的過去有一部分繼續被保存。未來世代可以看著這些畫，想像生活在甘蔗曾是台灣文化重要部分的年代是什麼樣子。那些日子也許早已遠去，但只要有幾個學生看到壁畫後好奇地提出問題，甘蔗與運送甘蔗的小火車的故事就會永遠流傳下去。",
    ],
    3: [
        "我小時候覺得刮鬍子一定很好玩。每當父親刮鬍子時，我都會站在洗手台旁看著。那時候看起來很令人興奮。有時我甚至會把肥皂抹在臉上，再用玩具刮鬍刀假裝刮掉。刮鬍子的樂趣早已消失。以前看起來很有趣的事，現在只是一件煩人的麻煩事。我應該每天刮鬍子，而且如果每天刮會輕鬆得多，但我通常會等到鬍渣癢到受不了才刮。如果我真的很懶，鬍子長得太長，就必須先用修鬍器再用刮鬍刀。修鬍器會把鬍鬚修短到可以刮的長度。省略這個步驟會讓刮鬍子有點痛。用刮鬍刀拖過長鬍子時，與其說是在切斷鬍子，不如說是在拉扯鬍子，完全不舒服。更不用說刀片很快就會被較長的鬍鬚堵住。基本上，我的刮鬍包裡有三樣東西：刮鬍凝膠、可充電修鬍器和刮鬍刀。雖然照片中顯示的是刮鬍泡，但我通常用凝膠，因為它能刮得更貼近也更順。除此之外，我只需要水和一面鏡子。我通常在洗澡後刮鬍子，因為熱水會讓皮膚毛孔張開，這能幫助刮得更乾淨，也讓刮鬍刀更容易滑過皮膚。如果沒有時間先洗澡，我至少會好好洗臉，去除髒汙和油脂，並用熱毛巾打開毛孔。",
        "使用刮鬍凝膠時，只需要在手掌上擠一點點，因為當你把凝膠抹在臉上時，它會膨脹並變白。凝膠也能作為刮鬍刀的潤滑劑，減少摩擦，讓刮鬍更舒服。使用刮鬍泡時則需要大方一些。如果用得太少，它會在臉上乾掉，你可能得重新塗一次。你應該只把凝膠或泡沫塗在需要刮的地方，以免浪費。對我來說，這差不多就是眼睛以下、胸口以上的所有地方。",
        "開始刮鬍子時，先把刮鬍刀放在流動的熱水中沖一下，但不要太熱，否則金屬刀片可能會燙傷你。接著，沿著鬍子生長的方向，以向下的動作慢慢拉動刮鬍刀。每刮一下都應該停下來沖洗刮鬍刀，避免刀片被鬍子堵住。下顎線和下巴周圍要小心。即使用現代安全刮鬍刀，也很容易刮傷自己。嘴唇邊緣特別脆弱，這是我刮鬍子時少數曾割傷自己的地方之一。把臉上發癢的小鬍鬚清除後，剩下的只是把臉沖乾淨並用毛巾擦乾。這時也可以擦鬍後水，但除非我大約一個月沒刮鬍子，否則我通常不會這麼做。我也建議不要在運動前刮鬍子。每次我試過這樣做，都會因為受刺激的毛孔被汗水覆蓋而起疹子。",
        "刮鬍子的最佳時間大概是晚上，空氣涼爽的時候。這也讓你的臉在再次暴露於汗水、油脂和陽光之前，有一整晚的休息時間。話雖如此，大多數男人還是把刮鬍子當成早晨例行事務的一部分。我想有些人比較喜歡乾淨刮好的樣子。",
    ],
    4: [
        "今天我看了一個覺得很有趣的電視節目。節目介紹一位來自台灣新化的男子，他因為一項非常特殊的熱情，在自己的社區中幾乎具有代表性地位。這份熱情最後帶他越過台灣海峽，前往中國大陸，也就是這項不尋常迷戀的發源地：鬥蟋蟀。",
        "鬥蟋蟀的根源可以追溯到唐朝，幾世紀以來已深深扎根於許多亞洲國家的文化土壤中。很久以前，它是皇室的消遣，但隨著時間過去，越來越多平民也參與其中。即使到了現代，並且出現一波波倡導善待所有生命的潮流，這項活動仍持續發展。",
        "謝絕安，或當地人稱呼他的 Ango，和他那個年代許多台灣孩子一樣，常和朋友參加鬥蟋蟀。隨著時間過去，許多孩子會長大並拋下童年嗜好，但 Ango 沒有。他對蟋蟀的熱情延續到成年，如今他已成為這個主題的權威，甚至協助全國各地的大學進行研究。",
        "新化之所以成為台灣鬥蟋蟀之都，要歸功於 Ango。他和幾位好友多年前在這個小鎮舉辦第一場正式鬥蟋蟀比賽。這項「運動」迅速流行起來，而年度賽事至今仍持續舉辦，並吸引越來越多來自全島的參賽者。",
        "Ango 作為認真且「不可小看」的蟋蟀訓練師的名聲，最後傳到了中國大陸。在那裡，世界最大比賽之一的主辦人邀請 Ango 參賽。中國方面想看看這個人到底有多厲害。",
        "Ango 非常慎重看待這項邀請。他花了數月繁殖、挑選並訓練最優秀、最兇猛的蟋蟀參加比賽。他不知道在鬥蟋蟀的發源地會有什麼等著他，但他盡了一切努力準備。",
        "當 Ango 終於抵達上海參賽時，他大為震撼。在他的故鄉台灣，鬥蟋蟀只是一種嗜好；但在這座世界最大城市之一裡，它卻是一門大生意。那裡有市場可以購買蟋蟀，以及飼養照顧蟋蟀所需的各種用品。他也注意到這項活動較陰暗的一面，因為他看到許多人重金下注在鬥蟋蟀比賽上。",
        "Ango 和主辦人的會面是非常吉利而重要的場合。兩人都對彼此的做法和行內技巧感到好奇。他們花了好幾個小時交換資訊並談論專業經驗。",
        "主辦人立刻感到驚訝的一點，是 Ango 的蟋蟀體型很大。和牠們的中國表親，也就是有時被稱為黑龍的蟋蟀相比，Ango 的蟋蟀非常巨大。雖然他對看到一種新的蟋蟀感到敬畏，但也開始擔心 Ango 的蟋蟀可能會摧毀他引以為傲、被他親切稱為「將軍」的蟋蟀。",
        "當兩隻蟋蟀和牠們的主人準備開戰時，空氣中充滿緊張氣氛。沒有人知道任何一隻蟋蟀會有什麼反應。那隻台灣大蟋蟀確實有令人印象深刻的體型優勢，但經過長途飛行後似乎有些無精打采。將軍雖然較小，卻屬於非常具攻擊性的品種，常把對手肢解，但牠從未面對過如此巨大的對手。",
        "兩隻蟋蟀終於被放進透明的鬥盒中。每位訓練師都用自己的技巧讓蟋蟀進入戰鬥狀態。最後，分隔兩者的門被拉起，牠們衝向彼此；但當將軍看到黑龍的體型時，牠改變主意並跳出籠子。黑龍勝利地鳴叫。整件事其實相當虎頭蛇尾，不過我想雙方都學到了很多。",
        "Ango 在中國停留後回到家鄉。雖然實際上並沒有真正的「戰鬥」，但他學到了許多傳統中國蟋蟀飼養與戰鬥的知識，並迫不及待想和朋友分享。他也從一個與自己故鄉關係常常不太好的地方，獲得了一位新朋友。這會是一段他可能永遠珍惜在心的經驗。",
    ],
    5: [
        "這是加拿大亞伯達省小奴湖附近的奴河。它是一條偏遠的河流，蜿蜒穿過亞伯達廣闊北方森林的中心。搭船在河上看到的許多地區只能從水路抵達，因此仍保持原始自然。除了偶爾有一塊空地上建著房子，稍微破壞原本如畫的景色之外，整體景觀幾乎未受干擾。",
        "第一張照片中，河的遠端有三、四艘獨木舟。船上的人順著水流漂行，尋找可能藏在河岸濃密樹葉中的野生動物。水鳥和鳴禽很常見，四周都看得到，或至少聽得到；但較大型的動物就不那麼容易被發現。",
        "這裡的森林非常茂密，動物比較可能先看到你，而不是你看到牠們。牠們也更可能在你抵達很久之前就聽到你，並在你瞥見牠們之前就走開。即便如此，划獨木舟的人仍然觀察著，希望看到一些動靜，因為那往往是唯一能暴露動物藏身處的線索。",
        "第二張照片中，我們看到一位英俊的年輕男子和他的朋友快樂地划著船順流而下。他們都穿著救生衣，因為他們不會游泳；但即使會游泳，划獨木舟時穿救生衣仍是好主意，因為河流的水流可能很強，也很難逆流游動。只有愚蠢的人才會不穿救生衣去挑戰命運。",
        "你可能會注意到划船的人在獨木舟兩側划槳。這是為了讓船直行。如果兩人都在同一側划，他們最後會像單臂游泳的人一樣原地繞圈。",
        "讓獨木舟朝正確方向前進聽起來可能很容易，但實際上比看起來困難。划槳者之間需要良好的溝通。坐在後面的人就像船長，通常會下達一些基本指令。他會告訴前面的人該在哪一側划、該用多大的力氣。接著他可以調整自己的划槳速度和動作，好讓船保持正確方向。",
        "這裡另一位男子划著皮艇靠到獨木舟旁。皮艇和獨木舟相似，但頂部是封閉的，只有一個圓形開口讓划皮艇的人坐進去。這種封閉設計讓皮艇翻覆時更容易翻正。只要你能憋氣且不驚慌，甚至不必離開船。只要簡單拉動槳，就能把皮艇翻回正位；但獨木舟則必須鑽到船下，把它從水中推起來並試著翻正。皮艇流線型的形狀也讓它比獨木舟更快、更靈活。",
        "皮艇使用的槳也和獨木舟槳不同：它有兩片槳葉，而不是一片。這表示你不必像划獨木舟那樣換邊，因為兩端都有槳葉。這種設計在水流更快、更湍急的地方非常有用，因為為了避免撞上岩石，瞬間改變方向是關鍵。",
    ],
}

EXTRA_VOCAB = {
    "sleepy": ("adj.", "寧靜的；冷清的"),
    "congested": ("adj.", "壅塞的"),
    "skyscraper": ("n.", "摩天大樓"),
    "unremarkable": ("adj.", "不起眼的"),
    "disproportionately": ("adv.", "不成比例地"),
    "sparsely": ("adv.", "稀疏地"),
    "unobstructed": ("adj.", "無阻礙的"),
    "minimalist": ("adj.", "極簡的"),
    "bewilder": ("v.", "使困惑"),
    "manicured": ("adj.", "修剪整齊的"),
    "serenity": ("n.", "寧靜"),
    "symmetrical": ("adj.", "對稱的"),
    "accentuate": ("v.", "凸顯"),
    "foyer": ("n.", "門廳"),
    "abbey": ("n.", "修道院"),
    "vegetarian": ("adj.", "素食的"),
    "delicacy": ("n.", "佳餚"),
    "serenade": ("v.", "唱歌或演奏給…聽"),
    "hospitality": ("n.", "款待"),
    "pessimist": ("n.", "悲觀者"),
    "mural": ("n.", "壁畫"),
    "portray": ("v.", "描繪"),
    "harvesting": ("n.", "採收"),
    "lush": ("adj.", "茂盛的"),
    "depiction": ("n.", "描繪"),
    "toil": ("v.", "辛苦工作"),
    "enormity": ("n.", "巨大規模"),
    "foreground": ("n.", "前景"),
    "preliminary": ("adj.", "初步的"),
    "mastermind": ("n.", "策劃者"),
    "ambitious": ("adj.", "有雄心的"),
    "invaluable": ("adj.", "非常寶貴的"),
    "facelift": ("n.", "整修；美化"),
    "shaving": ("n.", "刮鬍子"),
    "lather": ("v.", "塗上泡沫"),
    "stubble": ("n.", "鬍渣"),
    "trimmer": ("n.", "修剪器"),
    "whisker": ("n.", "鬍鬚"),
    "lubricant": ("n.", "潤滑劑"),
    "friction": ("n.", "摩擦"),
    "vulnerable": ("adj.", "脆弱的；易受傷的"),
    "aftershave": ("n.", "鬍後水"),
    "ritual": ("n.", "例行儀式"),
    "cricket": ("n.", "蟋蟀"),
    "iconic": ("adj.", "具代表性的"),
    "peculiar": ("adj.", "特殊的；奇特的"),
    "obsession": ("n.", "著迷"),
    "pastime": ("n.", "消遣"),
    "commoner": ("n.", "平民"),
    "authority": ("n.", "權威"),
    "reckon": ("v.", "認為；估計"),
    "breed": ("v.", "繁殖"),
    "fierce": ("adj.", "兇猛的"),
    "paraphernalia": ("n.", "用品；裝備"),
    "auspicious": ("adj.", "吉利的"),
    "annihilate": ("v.", "殲滅；徹底擊敗"),
    "lethargic": ("adj.", "無精打采的"),
    "dismember": ("v.", "肢解"),
    "anticlimactic": ("adj.", "虎頭蛇尾的"),
    "husbandry": ("n.", "飼養管理"),
    "canoe": ("n.", "獨木舟"),
    "kayak": ("n.", "皮艇"),
    "boreal": ("adj.", "北方的"),
    "blemish": ("v.", "破壞；玷污"),
    "picturesque": ("adj.", "如畫的"),
    "occupant": ("n.", "乘坐者"),
    "foliage": ("n.", "樹葉"),
    "riverbank": ("n.", "河岸"),
    "lifejacket": ("n.", "救生衣"),
    "paddle": ("v.", "划槳"),
    "capsize": ("v.", "翻覆"),
    "streamlined": ("adj.", "流線型的"),
    "maneuverable": ("adj.", "靈活可操控的"),
    "turbulent": ("adj.", "湍急的"),
    "split-second": ("adj.", "瞬間的"),
}

STOP = {
    "about", "after", "again", "also", "another", "because", "before", "between", "could",
    "every", "first", "from", "have", "into", "like", "little", "more", "most", "much",
    "only", "other", "over", "people", "picture", "really", "said", "some", "something",
    "that", "their", "them", "then", "there", "these", "they", "this", "those", "through",
    "time", "when", "where", "which", "while", "with", "would", "your",
}


def clean(text):
    return re.sub(r"\s+", " ", text.replace("\u00a0", " ")).strip()


def load_vocab_map():
    vocab = dict(EXTRA_VOCAB)
    for name in ("basic.json", "intermediate.json", "advanced.json", "description.json"):
        path = ROOT / "data" / name
        if not path.exists():
            continue
        data = json.loads(path.read_text(encoding="utf-8"))
        for units in data.values():
            for unit in units:
                for item in unit.get("vocab", []):
                    word = clean(item.get("w", "")).lower()
                    if word and item.get("zh"):
                        vocab.setdefault(word, (item.get("pos") or "", item["zh"]))
    return vocab


def parse_docx(path):
    texts = [clean(p.text) for p in Document(path).paragraphs]
    texts = [t for t in texts if t]
    title = re.sub(r"^Unit\s*\d+\s*:?\s*", "", texts[0], flags=re.I).strip()
    paras = []
    for text in texts[1:]:
        if text.lower().startswith("questions:"):
            break
        paras.append(text)
    return title, paras


def extract_images(path, unit):
    doc = Document(path)
    rels = [rel for rel in doc.part.rels.values() if "image" in rel.reltype]
    out = []
    for index, rel in enumerate(rels, 1):
        if rel.target_part.content_type != "image/jpeg":
            continue
        target = IMG_DIR / f"adv-b{BOOK:02d}-u{unit:02d}-{index}.jpg"
        target.write_bytes(rel.target_part.blob)
        out.append(f"/assets/img/booklets/{target.name}")
    return out


def make_vocab(paras, title, vocab_map):
    text = " ".join([title] + paras).lower()
    phrases = []
    for phrase in sorted(vocab_map, key=len, reverse=True):
        if " " in phrase and phrase in text:
            phrases.append(phrase)
    words = re.findall(r"[a-zA-Z][a-zA-Z'-]{3,}", text)
    counts = Counter(w.strip("'").lower() for w in words if w.lower() not in STOP)
    candidates = phrases + [word for word, _ in counts.most_common(120)]
    out = []
    used = set()
    for candidate in candidates:
        key = candidate.lower()
        if key in used or key not in vocab_map:
            continue
        pos, zh = vocab_map[key]
        out.append({"n": len(out) + 1, "w": candidate, "pos": pos, "zh": zh})
        used.add(key)
        if len(out) >= 30:
            break
    return out


def main():
    vocab_map = load_vocab_map()
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    units = []
    for unit, filename in UNIT_FILES.items():
        path = SOURCE_DIR / filename
        title, paras = parse_docx(path)
        units.append({
            "unit": unit,
            "title": title,
            "paras": paras,
            "questions": [],
            "vocab": make_vocab(paras, title, vocab_map),
            "translation": " ".join(TRANSLATIONS[unit]),
            "answers": [],
            "photos": extract_images(path, unit),
            "pdf": "",
            "audio": {},
        })
    data[str(BOOK)] = units
    DATA_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
