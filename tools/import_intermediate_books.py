#!/usr/bin/env python3
import json
import os
import re
import shutil
from collections import Counter
from pathlib import Path

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = Path.home() / "Downloads"
BOOK_FOLDERS = {
    5: DOWNLOADS / "中級閱讀 Book05",
    6: DOWNLOADS / "中級閱讀 Book06",
}
DATA_PATH = ROOT / "data" / "intermediate.json"
IMG_DIR = ROOT / "assets" / "img" / "booklets"

EXTRA_VOCAB = {
    "tangerine": ("n.", "柑橘"),
    "orchard": ("n.", "果園"),
    "foliage": ("n.", "葉子；樹葉"),
    "overwhelming": ("adj.", "壓倒性的；強烈的"),
    "daunting": ("adj.", "令人畏懼的"),
    "whine": ("v.", "抱怨；發牢騷"),
    "competitive": ("adj.", "好競爭的"),
    "tauntingly": ("adv.", "嘲弄地"),
    "relay": ("n.", "接力賽"),
    "baton": ("n.", "接力棒"),
    "false start": ("n.", "起跑犯規"),
    "penalize": ("v.", "處罰"),
    "disqualify": ("v.", "取消資格"),
    "stagger": ("v.", "使錯開"),
    "handoff": ("n.", "交棒"),
    "transition": ("n.", "轉換；交接"),
    "anchor": ("n.", "最後一棒選手"),
    "strategy": ("n.", "策略"),
    "trash": ("n.", "垃圾"),
    "garbage": ("n.", "垃圾"),
    "decompose": ("v.", "分解；腐爛"),
    "stench": ("n.", "惡臭"),
    "lid": ("n.", "蓋子"),
    "maggot": ("n.", "蛆"),
    "dumpster": ("n.", "大型垃圾箱"),
    "color blindness": ("n.", "色盲"),
    "color-blind": ("adj.", "色盲的"),
    "distinguish": ("v.", "辨別"),
    "shade": ("n.", "色調"),
    "handicap": ("n.", "障礙；不利條件"),
    "fruitful": ("adj.", "有成果的"),
    "taekwondo": ("n.", "跆拳道"),
    "demonstration": ("n.", "示範"),
    "martial art": ("n.", "武術"),
    "flexible": ("adj.", "柔軟的"),
    "cynic": ("n.", "憤世嫉俗者"),
    "graduation": ("n.", "畢業典禮"),
    "ceremony": ("n.", "典禮"),
    "routine": ("n.", "例行表演"),
    "audience": ("n.", "觀眾"),
    "mingle": ("v.", "交際；混在一起"),
    "chalkboard": ("n.", "黑板"),
    "eraser": ("n.", "板擦"),
    "dusty": ("adj.", "滿是灰塵的"),
    "retrieve": ("v.", "取回"),
    "intention": ("n.", "意圖"),
    "upset": ("adj.", "難過的；生氣的"),
    "sympathetic": ("adj.", "有同情心的"),
    "checkup": ("n.", "健康檢查"),
    "nurse": ("n.", "護士"),
    "eyesight": ("n.", "視力"),
    "weigh": ("v.", "秤重"),
    "blown away": ("ph.", "被風吹走"),
    "typhoon": ("n.", "颱風"),
    "flashlight": ("n.", "手電筒"),
    "curtain": ("n.", "窗簾"),
    "sunlight": ("n.", "陽光"),
    "repaint": ("v.", "重新油漆"),
    "performer": ("n.", "表演者"),
    "erhu": ("n.", "二胡"),
    "instrument": ("n.", "樂器"),
    "library": ("n.", "圖書館"),
    "collection": ("n.", "收藏"),
    "disappearing": ("adj.", "消失的"),
    "pen": ("n.", "筆"),
    "yoga": ("n.", "瑜伽"),
    "bookstore": ("n.", "書店"),
    "halloween": ("n.", "萬聖節"),
    "costume": ("n.", "服裝"),
    "birthday": ("n.", "生日"),
    "tricycle": ("n.", "三輪車"),
    "tee ball": ("n.", "樂樂棒球"),
    "gps": ("n.", "衛星導航"),
    "satellite": ("n.", "衛星"),
    "destination": ("n.", "目的地"),
    "altitude": ("n.", "海拔高度"),
    "coordinate": ("n.", "座標"),
}

STOP = {
    "about", "after", "again", "also", "another", "because", "before", "between", "could",
    "every", "first", "from", "have", "into", "like", "little", "more", "most", "much",
    "only", "other", "over", "people", "picture", "really", "said", "some", "something",
    "that", "their", "them", "then", "there", "these", "they", "this", "those", "through",
    "time", "when", "where", "which", "while", "with", "would", "your",
}

QUESTION_BANK = {
    (5, 1): [
        ("What fruit does Rick help his grandmother pick?", "He helps his grandmother pick tangerines."),
        ("Why does the task seem daunting to Rick?", "There are so many tangerines that it feels like they will never finish."),
        ("How does Rick's grandmother make the work more fun?", "She times herself filling boxes and tries to beat her own time."),
    ],
    (5, 2): [
        ("What are the girls waiting for at the start of the race?", "They are waiting for the starting pistol to fire."),
        ("Why are the starting lines staggered around the bend?", "They are staggered so each lane is the same distance from the finish line."),
        ("Why is teamwork important in a relay race?", "A team can lose if it drops the baton or makes a poor handoff."),
    ],
    (5, 3): [
        ("What does Tony almost forget to take with the garbage?", "He almost forgets to take the recycling."),
        ("Why is Tony in a hurry?", "He is trying to catch the garbage truck before it leaves."),
        ("What does Tony learn from taking out the trash?", "He learns that garbage and recycling have to be handled carefully and on time."),
    ],
    (5, 4): [
        ("Why does Eric think something is wrong with his eyesight?", "He sees a shirt as blue when others say it is green."),
        ("What test does the nurse use to check Eric's color vision?", "She uses cards with colored dots and hidden numbers or letters."),
        ("What kind of color problem does Eric have?", "He has a slight blue-green deficiency."),
    ],
    (5, 5): [
        ("Why did John first want to learn martial arts?", "He wanted to learn after watching old Bruce Lee movies."),
        ("Why did John and his friends lose interest in Judo?", "They thought it was too much like wrestling and did not have enough kicking."),
        ("What impressed the crowd during the Taekwondo demonstration?", "The boys' kicks, movements, and performance impressed the crowd."),
    ],
    (5, 6): [
        ("Where does the graduation ceremony take place?", "It takes place at Jhutang Elementary School."),
        ("What does the principal give to each student?", "The principal gives each student a diploma rolled like a scroll."),
        ("What do the students do after the formal ceremony?", "They mingle with the crowd and take pictures with friends and family."),
    ],
    (5, 7): [
        ("Why are chalkboards and whiteboards useful for teachers?", "They help teachers show words, ideas, and drawings to the whole class."),
        ("Why should chalkboard erasers be cleaned regularly?", "They become dirty with chalk dust and stop working well."),
        ("Why do many computer labs use whiteboards instead of chalkboards?", "Whiteboards create less dust that could harm sensitive computer parts."),
    ],
    (5, 8): [
        ("Why is James's back sore?", "His back is sore because he sat hunched over all day finishing a report."),
        ("What does Brett try to do for James?", "He tries to help James even though the help causes trouble."),
        ("Why are Brett's intentions called good?", "He means to help his brother, even if the result is not what James wanted."),
    ],
    (5, 9): [
        ("What happens during the annual health checkup?", "A nurse gives the students a physical examination."),
        ("Why do many students not mind the checkup?", "It takes the whole morning, so they do not have regular classes."),
        ("What does the narrator joke about at the end?", "He jokes that his height is the problem, not his weight."),
    ],
    (5, 10): [
        ("Why is Rachel worried at the beginning of the story?", "She hears that a typhoon is coming."),
        ("What does Rachel realize about her emergency food supply?", "She realizes she ate all of it during a late-night study session."),
        ("What happens when Rachel steps outside?", "The strong wind blows her away."),
    ],
    (6, 1): [
        ("Why does Luke pull over to the side of the road?", "He pulls over because he is lost again."),
        ("What does Luke use to find his destination?", "He uses a GPS to find his destination."),
        ("What information can a GPS provide?", "It can provide directions, altitude, speed, ETA, and coordinates."),
    ],
    (6, 2): [
        ("Where is the man performing?", "He is performing outside a temple."),
        ("What instrument is the man playing?", "He is playing an erhu."),
        ("How do passersby react to the performance?", "Some applaud or smile, while others frown or show little interest."),
    ],
    (6, 3): [
        ("What unusual thing do Guy and Jacque see?", "They see a mobile library parked near the temple."),
        ("Why is the library called a library on wheels?", "It brings books to people in a vehicle."),
        ("What do the boys decide to do in the end?", "They decide to go inside and look at books."),
    ],
    (6, 4): [
        ("Why is Julie late?", "She says she had to drop off her Doraemon collection at the babysitter."),
        ("What is Julie worried about?", "She worries that a cat burglar might steal her collection."),
        ("How does Susan react to Julie's explanation?", "Susan thinks Julie's concern is absurd."),
    ],
    (6, 5): [
        ("What problem does Wendy often have?", "She is always losing her pens and pencils."),
        ("How do people search for Wendy's pen?", "They search around her desk and then around the classroom."),
        ("Why is the situation funny?", "The missing pens become a mystery even though Wendy is otherwise a very good girl."),
    ],
    (6, 6): [
        ("What word does Robert misunderstand?", "He mistakes yoga for yogurt."),
        ("Where do Henrik and Robert join the class?", "They join the yoga class in the school activity room."),
        ("What embarrassing thing happens to Robert?", "His pants tear while he tries to do a yoga pose."),
    ],
    (6, 7): [
        ("Where is the book section described in the passage?", "It is on the second floor of RT-Mart."),
        ("Why does the narrator often go there?", "The narrator is a bookworm and likes looking for new books to read."),
        ("Why does the narrator compare bookstores and libraries?", "Some people read books in stores without buying them, as if the store were a library."),
    ],
    (6, 8): [
        ("What kinds of creatures are associated with Halloween?", "Ghosts, vampires, skeletons, witches, mummies, and zombies are associated with Halloween."),
        ("Why do children enjoy Halloween?", "They enjoy costumes, candy, and the excitement of the holiday."),
        ("Who has to deal with the real nightmare after Halloween?", "Parents, teachers, and especially dentists have to deal with the problems caused by too much candy."),
    ],
    (6, 9): [
        ("What present does Trina receive for her birthday?", "She receives a new bicycle."),
        ("Why is Trina so happy?", "She had longed for the bicycle and thinks it is the best present ever."),
        ("What will Trina and her bicycle do together?", "They will ride around the neighborhood and explore the world around them."),
    ],
    (6, 10): [
        ("Why is tee ball useful for young children?", "It lets children practice baseball before they can hit a moving ball well."),
        ("How does the tee help the hitter?", "It holds the ball in place at the correct height."),
        ("What do children learn by playing tee ball?", "They develop basic skills and learn the rules of baseball."),
    ],
}


def clean(s):
    return re.sub(r"\s+", " ", s.replace("\u00a0", " ")).strip()


def is_chinese(s):
    return bool(re.search(r"[\u4e00-\u9fff]", s))


def load_vocab_map():
    vocab = dict(EXTRA_VOCAB)
    for name in ("basic.json", "intermediate.json", "advanced.json", "description.json"):
        path = ROOT / "data" / name
        if not path.exists():
            continue
        data = json.loads(path.read_text(encoding="utf-8"))
        for units in data.values():
            for unit in units:
                for item in unit.get("vocab", []):
                    w = clean(item.get("w", "")).lower()
                    if w and item.get("zh"):
                        vocab.setdefault(w, (item.get("pos") or "", item["zh"]))
    return vocab


def pick_docx(folder, book, unit):
    matches = sorted(folder.glob(f"*Book{book:02d}_Unit{unit:02d}*.docx"))
    if not matches:
        raise FileNotFoundError(f"Missing docx for book {book} unit {unit}")
    def zh_count(path):
        return sum(1 for p in Document(path).paragraphs if is_chinese(p.text))
    return max(matches, key=zh_count)


def parse_docx(path):
    texts = [clean(p.text) for p in Document(path).paragraphs]
    texts = [t for t in texts if t]
    title_line = texts[0]
    title = re.sub(r"^Unit\s*\d+\s*:?\s*", "", title_line, flags=re.I).strip()
    body = texts[1:]
    first_zh = next((i for i, t in enumerate(body) if is_chinese(t)), len(body))
    paras = body[:first_zh]
    zh = body[first_zh:]
    if zh and re.match(r"^第.+課", zh[0]):
        zh = zh[1:]
    return title, paras, " ".join(zh)


def extract_images(pdf_path, book, unit):
    doc = fitz.open(pdf_path)
    seen = set()
    photos = []
    for page in doc:
        for img in page.get_images(full=True):
            xref = img[0]
            if xref in seen:
                continue
            seen.add(xref)
            base = doc.extract_image(xref)
            width, height = base["width"], base["height"]
            ext = "jpg" if base["ext"].lower() in {"jpeg", "jpg"} else base["ext"].lower()
            if ext != "jpg" or width < 250 or height < 250:
                continue
            photos.append(base["image"])
    rels = []
    for i, blob in enumerate(photos, 1):
        out = IMG_DIR / f"inter-b{book:02d}-u{unit:02d}-{i}.jpg"
        out.write_bytes(blob)
        rels.append(f"/assets/img/booklets/{out.name}")
    return rels


def make_vocab(paras, title, vocab_map):
    text = " ".join([title] + paras).lower()
    phrases = []
    for phrase in sorted(vocab_map, key=len, reverse=True):
        if " " in phrase and phrase in text:
            phrases.append(phrase)
    words = re.findall(r"[a-zA-Z][a-zA-Z'-]{3,}", text)
    counts = Counter(w.strip("'").lower() for w in words if w.lower() not in STOP)
    candidates = phrases + [w for w, _ in counts.most_common(80)]
    out = []
    used = set()
    for cand in candidates:
        key = cand.lower()
        if key in used or key not in vocab_map:
            continue
        pos, zh = vocab_map[key]
        out.append({"n": len(out) + 1, "w": cand, "pos": pos, "zh": zh})
        used.add(key)
        if len(out) >= 10:
            break
    return out


def fallback_qa(title, paras):
    who = title.split()[0].strip("'’") if title else "the person"
    topic = title[0].lower() + title[1:] if title else "the story"
    return [
        (f"What is this unit mainly about?", f"It is mainly about {topic}."),
        (f"What problem or task does {who} face in the story?", "The story describes the main task or situation and how it develops."),
        ("What can readers learn from this passage?", "They can learn useful English words and details about the situation in the passage."),
    ]


def main():
    vocab_map = load_vocab_map()
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    for book, folder in BOOK_FOLDERS.items():
        units = []
        for unit in range(1, 11):
            docx_path = pick_docx(folder, book, unit)
            pdf_path = folder / f"Copy of Intermediate_Book{book:02d}_Unit{unit:02d}.pdf"
            title, paras, translation = parse_docx(docx_path)
            qa = QUESTION_BANK.get((book, unit), fallback_qa(title, paras))
            units.append({
                "unit": unit,
                "title": title,
                "paras": paras,
                "questions": [q for q, _ in qa],
                "vocab": make_vocab(paras, title, vocab_map),
                "translation": translation,
                "answers": [a for _, a in qa],
                "photos": extract_images(pdf_path, book, unit),
                "pdf": f"Intermediate_Book{book:02d}_Unit{unit:02d}.pdf",
                "audio": {},
            })
        data[str(book)] = units
    DATA_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    release_dir = ROOT / "_release_assets" / "intermediate-pdf"
    release_dir.mkdir(parents=True, exist_ok=True)
    for book, folder in BOOK_FOLDERS.items():
        for unit in range(1, 11):
            src = folder / f"Copy of Intermediate_Book{book:02d}_Unit{unit:02d}.pdf"
            dst = release_dir / f"Intermediate_Book{book:02d}_Unit{unit:02d}.pdf"
            shutil.copy2(src, dst)


if __name__ == "__main__":
    main()
